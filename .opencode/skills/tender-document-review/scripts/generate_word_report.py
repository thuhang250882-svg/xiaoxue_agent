#!/usr/bin/env python3
"""Generate a formal Chinese DOCX report from validated tender-review JSON."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from validate_review_output import validate


def font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def section_format(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Cm(29.7 if landscape else 21)
    section.page_height = Cm(21 if landscape else 29.7)
    section.top_margin = Cm(2.2 if landscape else 3.7)
    section.bottom_margin = Cm(2.2 if landscape else 3.5)
    section.left_margin = Cm(2 if landscape else 2.8)
    section.right_margin = Cm(2 if landscape else 2.6)


def page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(paragraph.add_run("— "), "宋体", 10.5)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    cached = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    cached.append(text)
    field.append(cached)
    paragraph._p.append(field)
    font(paragraph.add_run(" —"), "宋体", 10.5)


def styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "仿宋"
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.first_line_indent = Pt(32)
    for name, family, size, bold in (
        ("Title", "宋体", 22, True),
        ("Heading 1", "黑体", 16, False),
        ("Heading 2", "楷体", 16, True),
    ):
        style = document.styles[name]
        style.font.name = family
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), family)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True


def heading(document: Document, number: str, title: str) -> None:
    document.add_paragraph(f"{number}、{title}", style="Heading 1")


def body(document: Document, text: str, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(32) if indent else Pt(0)
    font(paragraph.add_run(text), "仿宋", 16)


def cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    font(paragraph.add_run(text or "—"), "宋体", 10.5, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell) -> None:
    value = OxmlElement("w:shd")
    value.set(qn("w:fill"), "D9D9D9")
    cell._tc.get_or_add_tcPr().append(value)


def repeat_header(row) -> None:
    value = OxmlElement("w:tblHeader")
    value.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(value)


def table(document: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    result = document.add_table(rows=1, cols=len(headers))
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.autofit = False
    for cell, value in zip(result.rows[0].cells, headers):
        cell_text(cell, value, True, True)
        shade(cell)
    repeat_header(result.rows[0])
    for row in rows:
        cells = result.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, row)):
            cell_text(cell, str(value), center=index in {0, 1})


def evidence(values: list[dict[str, Any]]) -> str:
    return "；".join(f"{item['file']}（{item['location']}）：{item['quote']}" for item in values) or "未提供"


def create(payload: dict[str, Any]) -> Document:
    document = Document()
    styles(document)
    section_format(document.sections[0])
    page_number(document.sections[0])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    font(title.add_run(f"{payload['metadata']['project_name']}\n投标文件审核报告"), "宋体", 22, True)
    table(
        document,
        ("项目", "内容"),
        [
            ("项目名称", payload["metadata"]["project_name"]),
            ("项目编号", payload["metadata"].get("project_number") or "待确认"),
            ("审核日期", payload["metadata"]["review_date"]),
            ("编制部门", payload["metadata"].get("department") or "待填写"),
        ],
    )
    document.add_page_break()

    heading(document, "一", "审核基本情况")
    table(
        document,
        ("序号", "角色", "文件", "提取方式", "覆盖范围"),
        [(str(index), item["role"], item["name"], item["extraction"], item["coverage"]) for index, item in enumerate(payload["metadata"]["files"], 1)],
    )

    heading(document, "二", "审核结论")
    summary = payload["summary"]
    body(document, f"审核建议：{summary['overall']}。")
    body(document, "响应统计：" + "，".join(f"{key}{value}项" for key, value in summary["status_counts"].items()) + "。")
    body(document, "风险统计：" + "，".join(f"{key}{value}项" for key, value in summary["risk_counts"].items()) + "。")
    body(document, "优先工作：" + ("；".join(summary["top_actions"]) if summary["top_actions"] else "暂无。"))

    heading(document, "三", "需领导决策/协调事项")
    if not payload["decision_items"]:
        body(document, "经本次审核，未识别出需领导决策/协调事项。")
    for item in payload["decision_items"]:
        body(document, f"{item['id']}  {item['decision']}。可选方案：{'；'.join(item['options'])}。", False)
        body(document, f"证据与影响：{item['evidence_and_impact']} 责任部门：{item['owner']}；最迟决策时间：{item['deadline']}。")

    heading(document, "四", "致命及高风险事项")
    high = [item for item in payload["items"] if item["risk_level"] in {"致命", "高"}]
    if not high:
        body(document, "本次审核未识别出致命或高风险事项；仍需完成报告所列人工核验。")
    else:
        table(document, ("编号", "类别", "风险", "问题", "整改动作"), [(item["id"], item["category"], item["risk_level"], item["finding"], item["remediation"]) for item in high])

    matrix = document.add_section(WD_SECTION.NEW_PAGE)
    section_format(matrix, True)
    page_number(matrix)
    heading(document, "五", "逐项审核情况")
    table(
        document,
        ("编号", "类别", "招标要求及证据", "投标响应及证据", "状态", "风险", "问题", "整改/责任/时限"),
        [
            (
                item["id"],
                item["category"],
                f"{item['requirement']}\n{evidence(item['tender_evidence'])}",
                evidence(item["bid_evidence"]),
                item["status"],
                item["risk_level"],
                item["finding"],
                f"{item['remediation']}\n责任：{item['owner']}；时限：{item['deadline']}",
            )
            for item in payload["items"]
        ],
    )

    portrait = document.add_section(WD_SECTION.NEW_PAGE)
    section_format(portrait)
    page_number(portrait)
    heading(document, "六", "评分差距与价格核算")
    body(document, payload.get("scoring_and_price") or "现有底稿未单列可确认的评分差距或价格核算结果，待资料完整后复核。")
    heading(document, "七", "交叉矛盾和偏差")
    conflicts = [item for item in payload["items"] if item["status"] in {"不符合", "部分符合"}]
    body(document, "；".join(f"{item['id']}：{item['finding']}" for item in conflicts) or "本次审核未单独识别出交叉矛盾或偏差事项。")
    heading(document, "八", "整改工作安排")
    actions = [item for item in payload["items"] if item["status"] not in {"符合", "不适用"}]
    if actions:
        table(document, ("编号", "整改事项", "风险", "责任人/部门", "时限"), [(item["id"], item["remediation"], item["risk_level"], item["owner"], item["deadline"]) for item in actions])
    else:
        body(document, "本次审核未形成新增整改事项。")
    heading(document, "九", "人工核验事项")
    for index, item in enumerate(payload["manual_checks"] or ["本次审核未列出额外人工核验事项。"], 1):
        body(document, f"{index}. {item}", False)
    heading(document, "十", "审核限制和未审范围")
    for index, item in enumerate(payload["limitations"] or ["本次审核未列出额外限制。"], 1):
        body(document, f"{index}. {item}", False)
    body(document, "本报告是基于已提供材料形成的辅助审核意见，不代替采购人、评标委员会、律师、监管机构或单位授权审批人的最终判断。")
    return document


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("json_file", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    payload = json.loads(args.json_file.read_text(encoding="utf-8-sig"))
    errors, warnings = validate(payload)
    for warning in warnings:
        print(f"WARNING: {warning}", file=sys.stderr)
    if errors:
        raise SystemExit("JSON 校验失败：\n" + "\n".join(f"- {error}" for error in errors))
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    create(payload).save(args.output_docx)
    print(f"Word report generated: {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
